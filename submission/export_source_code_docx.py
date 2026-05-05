"""Export a readable DOCX bundle of the repository's source code.

This is intended for academic submission requirements where a Word document
containing the project code is requested in addition to the raw repository.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = REPO_ROOT / "submission" / "source_code_bundle.docx"

ROOT_ITEMS = [
    "app.py",
    "main_crew.py",
    "generate_graphs.py",
    "render_mermaid_diagrams.mjs",
    "pyproject.toml",
    ".env.example",
    "agents",
    "helpers",
    "integrations",
    "routes",
    "templates",
    "static",
    "experiments",
    "tests",
    "dissertation_material/datapipeline_code.py",
]

INCLUDED_EXTENSIONS = {
    ".py",
    ".js",
    ".mjs",
    ".html",
    ".css",
    ".toml",
    ".json",
    ".yml",
    ".yaml",
    ".ini",
    ".sql",
    ".nix",
    ".txt",
}

EXCLUDED_FILES = {
    Path("static/output/out.txt"),
}


def iter_source_files() -> list[Path]:
    files: list[Path] = []
    for item in ROOT_ITEMS:
        path = REPO_ROOT / item
        if path.is_file():
            if path.suffix.lower() in INCLUDED_EXTENSIONS or path.name == ".env.example":
                files.append(path)
            continue
        if not path.is_dir():
            continue
        for child in sorted(path.rglob("*")):
            if not child.is_file():
                continue
            rel = child.relative_to(REPO_ROOT)
            if "__pycache__" in rel.parts:
                continue
            if rel in EXCLUDED_FILES:
                continue
            if child.suffix.lower() not in INCLUDED_EXTENSIONS:
                continue
            files.append(child)
    deduped = sorted({file.relative_to(REPO_ROOT): file for file in files}.values(), key=lambda p: p.relative_to(REPO_ROOT).as_posix())
    return deduped


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)


def add_cover_page(document: Document, files: list[Path]) -> None:
    title = document.add_paragraph()
    run = title.add_run("Source Code Bundle")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.add_run("Automating Year-End Accounting Notes").bold = True
    subtitle.style = document.styles["Normal"]

    meta = document.add_paragraph()
    meta.add_run(
        "Readable submission export of repository source files.\n"
        f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}\n"
        f"Included files: {len(files)}"
    )

    document.add_paragraph(
        "This document contains the project source code in a single Word file for submission review. "
        "Live credentials are intentionally excluded."
    )

    groups = {
        "Application": ["app.py", "main_crew.py", "agents/", "routes/", "helpers/", "integrations/"],
        "Frontend": ["templates/", "static/"],
        "Experiments": ["experiments/", "generate_graphs.py", "render_mermaid_diagrams.mjs"],
        "Submission pipeline": ["dissertation_material/datapipeline_code.py", "tests/", "pyproject.toml", ".env.example"],
    }
    for section_name, included in groups.items():
        para = document.add_paragraph()
        run = para.add_run(f"{section_name}: ")
        run.bold = True
        para.add_run(", ".join(included))

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_file_heading(document: Document, relative_path: str, line_count: int) -> None:
    heading = document.add_paragraph()
    run = heading.add_run(relative_path)
    run.bold = True
    run.font.size = Pt(14)

    details = document.add_paragraph()
    details_run = details.add_run(f"{line_count} lines")
    details_run.italic = True
    details_run.font.size = Pt(9)


def add_code_block(document: Document, content: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(content)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def export_bundle() -> Path:
    files = iter_source_files()
    document = Document()
    configure_document(document)
    add_cover_page(document, files)

    for index, file_path in enumerate(files):
        relative = file_path.relative_to(REPO_ROOT).as_posix()
        text = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()
        numbered = "\n".join(f"{line_no:04d}: {line}" for line_no, line in enumerate(text, start=1))
        add_file_heading(document, relative, len(text))
        add_code_block(document, numbered if numbered else "0001: ")
        if index != len(files) - 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(export_bundle())
